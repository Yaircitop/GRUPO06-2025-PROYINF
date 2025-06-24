from django.shortcuts import render, redirect
from .forms import NewsForm, KeywordForm
from .models import News
from .utils import summarize_news
from django.db.models import Q
from django.http import HttpResponse
from docx import Document
from io import BytesIO

def handle_news_upload(request):
    news_form = NewsForm(request.POST)
    if news_form.is_valid():
        news_form.save()
        return redirect('crear_boletin')
    return news_form

def build_news_query(keywords):
    keyword_list = [word.strip() for word in keywords.split(",")]
    query = Q()
    for keyword in keyword_list:
        query |= Q(title__icontains=keyword) | Q(content__icontains=keyword)
    return query

def generate_docx(summarized_news):
    doc = Document()
    doc.add_heading("Boletín Informativo", level=1)
    doc.add_paragraph("")  # Espacio

    for item in summarized_news:
        doc.add_heading(item['title'], level=2)
        doc.add_paragraph(item['summary'])
        doc.add_paragraph("")  

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def crear_boletin(request):
    news_form = NewsForm()
    keyword_form = KeywordForm()
    summarized_news = []

    if request.method == "POST":
        if "upload_news" in request.POST:
            result = handle_news_upload(request)
            if isinstance(result, HttpResponse):
                return result
            news_form = result

        elif "search_and_summarize" in request.POST:
            keyword_form = KeywordForm(request.POST)
            if keyword_form.is_valid():
                query = build_news_query(keyword_form.cleaned_data['keywords'])
                relevant_news = News.objects.filter(query).distinct()
                summarized_news = summarize_news(relevant_news)
                buffer = generate_docx(summarized_news)

                response = HttpResponse(
                    buffer.getvalue(),
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = 'attachment; filename="boletin_resumen.docx"'
                return response

    return render(request, 'crear_boletin.html', {
        'news_form': news_form,
        'keyword_form': keyword_form,
        'summarized_news': summarized_news,
    })
